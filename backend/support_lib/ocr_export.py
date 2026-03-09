from docx import Document
from docx.shared import Pt
from pathlib import Path

def save_ocr_word(ocr_data, output_docx_path: Path):
    doc = Document()
    doc.add_heading(ocr_data["image"], level=1)

    for item in ocr_data["items"]:
        p = doc.add_paragraph(item["text"])
        p.paragraph_format.space_after = Pt(6)

    doc.save(output_docx_path)


import json

def save_ocr_json(ocr_data, output_path):
    """
    Save OCR result to JSON (safe for numpy types)
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    safe_data = to_builtin(ocr_data)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(safe_data, f, ensure_ascii=False, indent=2)

def to_builtin(obj):
    """
    Convert numpy / torch types to Python built-in types
    (safe for JSON serialization)
    """
    import numpy as np

    if isinstance(obj, dict):
        return {k: to_builtin(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [to_builtin(v) for v in obj]
    elif isinstance(obj, tuple):
        return [to_builtin(v) for v in obj]
    elif isinstance(obj, np.generic):
        return obj.item()   # numpy scalar → int / float
    else:
        return obj
    

from docx import Document
def save_analyzed_word(analyzed_data, output_path):
    doc = Document()
    doc.add_heading(
        f"Page: {analyzed_data.get('page','')}",
        level=2
    )
    for block in analyzed_data["blocks"]:
        # Paragraph
        if block["type"] == "paragraph":

            doc.add_paragraph(block["text"])

        # Table
        elif block["type"] == "table":
            rows = len(block["rows"])
            cols = len(block["rows"][0])
            table = doc.add_table(rows=rows, cols=cols)
            for r in range(rows):
                for c in range(cols):
                    table.cell(r, c).text = block["rows"][r][c]
            doc.add_paragraph("")
    doc.save(output_path)